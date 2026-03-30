"""DOCX document parser using markitdown."""

import io
import logging
from typing import Optional

from agent_rag.ingestion.parsing.base import BaseParser, ParsedDocument, ParsedImage
from agent_rag.ingestion.parsing.utils import (
    extract_title_from_text,
    normalize_text,
)

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    Parser for Microsoft Word .docx files.

    Uses markitdown library for robust DOCX to markdown conversion.

    Supports:
    - .docx, .doc files
    - application/vnd.openxmlformats-officedocument.wordprocessingml.document

    Priority: 10 (default for office documents)
    """

    def supports(self, source_type: str, extension: str, mime_type: str) -> bool:
        """Support .docx and .doc files."""
        if extension.lower() in ["docx", "doc"]:
            return True
        if mime_type and ("wordprocessingml" in mime_type or "msword" in mime_type):
            return True
        return False

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX file using markitdown.

        Args:
            content: Raw DOCX bytes
            filename: Original filename

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ImportError: If markitdown not installed
            Exception: If parsing fails
        """
        try:
            from markitdown import MarkItDown
        except ImportError:
            raise RuntimeError(
                "markitdown not installed. Run: pip install markitdown"
            )

        # Create temporary file-like object
        file_stream = io.BytesIO(content)

        # Convert to markdown
        md = MarkItDown()

        try:
            result = md.convert_stream(file_stream, file_extension=".docx")
        except Exception as e:
            logger.error(f"markitdown conversion failed: {e}")
            # Fallback to basic extraction if available
            return self._fallback_parse(content, filename)

        # Extract text
        text = result.text_content if hasattr(result, "text_content") else str(result)

        # Normalize text
        text = normalize_text(text)

        # Extract metadata
        metadata = {
            "filename": filename,
            "source_type": "file",
            "format": "docx",
        }

        # Try to extract title
        title = extract_title_from_text(text)
        if title:
            metadata["title"] = title

        # markitdown may provide additional metadata
        if hasattr(result, "metadata") and result.metadata:
            metadata.update(result.metadata)

        # Extract images (markitdown may include image references)
        images = self._extract_images(content, filename)

        return ParsedDocument(
            text=text,
            metadata=metadata,
            images=images,
            links=[],  # Could extract from markdown links
            tables=[],  # markitdown converts tables to markdown
        )

    def _fallback_parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Fallback parser using python-docx.

        Args:
            content: Raw DOCX bytes
            filename: Original filename

        Returns:
            ParsedDocument with basic text extraction
        """
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx not installed. Run: pip install python-docx"
            )

        try:
            # Load document
            file_stream = io.BytesIO(content)
            doc = Document(file_stream)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            text = "\n\n".join(paragraphs)
            text = normalize_text(text)

            # Basic metadata
            metadata = {
                "filename": filename,
                "source_type": "file",
                "format": "docx",
                "parser": "fallback_python_docx",
            }

            # Extract core properties if available
            if hasattr(doc, "core_properties"):
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.created:
                    metadata["created"] = props.created.isoformat()
                if props.modified:
                    metadata["modified"] = props.modified.isoformat()

            return ParsedDocument(
                text=text,
                metadata=metadata,
                images=[],
                links=[],
                tables=[],
            )

        except Exception as e:
            logger.error(f"Fallback DOCX parsing failed: {e}")
            raise RuntimeError(f"Failed to parse DOCX file: {e}")

    def _extract_images(
        self,
        content: bytes,
        filename: str,
    ) -> list[ParsedImage]:
        """
        Extract images from DOCX file.

        Args:
            content: Raw DOCX bytes
            filename: Original filename

        Returns:
            List of ParsedImage objects
        """
        images = []

        try:
            from docx import Document
            from docx.oxml import parse_xml
        except ImportError:
            return images

        try:
            file_stream = io.BytesIO(content)
            doc = Document(file_stream)

            # Extract inline images
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob

                        # Determine MIME type
                        content_type = image_part.content_type or "image/png"

                        # Determine extension
                        ext = content_type.split("/")[-1]
                        if ext == "jpeg":
                            ext = "jpg"

                        # Create ParsedImage
                        image_id = f"img_{len(images) + 1}"
                        images.append(
                            ParsedImage(
                                image_id=image_id,
                                content=image_bytes,
                                mime_type=content_type,
                                page_number=None,  # DOCX doesn't have page numbers
                                caption=None,
                            )
                        )

                    except Exception as e:
                        logger.warning(f"Failed to extract image from DOCX: {e}")
                        continue

        except Exception as e:
            logger.warning(f"Image extraction from DOCX failed: {e}")

        return images

    @property
    def priority(self) -> int:
        """Default priority for office documents."""
        return 10
